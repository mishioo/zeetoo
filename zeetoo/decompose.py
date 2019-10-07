import argparse
import logging
import olefile
import pathlib
import zipfile
try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

REL = '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'
OLEREL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject'
RELID = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id'
OLEOBJ = '{urn:schemas-microsoft-com:office:office}OLEObject'
WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PAR = WORD_NAMESPACE + 'p'
TXT = WORD_NAMESPACE + 't'

extensions = {
    'ChemDraw': 'cdx',
    'MestReNova': 'mnova'
}


def print_rec(element, depth=0, tab='|   '):
    for child in element.iterchildren():
        print(tab * depth + str(child))
        print_rec(child, depth + 1, tab)


def print_cls(element, depth=0, tab='|   '):
    for child in element.iterchildren():
        print(tab * depth + str(child.__class__))
        print_rec(child, depth + 1, tab)


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def copy_font(source, dest):
    attrs = [
        'all_caps', 'bold', 'complex_script', 'cs_bold', 'cs_italic',
        'double_strike', 'emboss', 'hidden', 'highlight_color',
        'imprint', 'italic', 'math', 'name', 'no_proof', 'outline',
        'rtl', 'shadow', 'size', 'small_caps', 'snap_to_grid', 'spec_vanish',
        'strike', 'subscript', 'superscript', 'underline', 'web_hidden'
    ]
    for attr in attrs:
        setattr(dest, attr, getattr(source, attr))
    if source.color.theme_color is not None:
        dest.color.theme_color = source.color.theme_color
    if source.color.rgb is not None:
        dest.color.rgb = source.color.rgb


def write_to_paragraph(source: Paragraph, dest: Paragraph):
    for run in source.runs:
        new = dest.add_run(run.text, run.style.name)
        copy_font(run.font, new.font)
        new.bold = run.bold
        new.italic = run.italic
        new.underline = run.underline
    # TODO: add math objects
    return dest


def copy_par_format(source, dest):
    attrs = [
        'alignment', 'first_line_indent', 'keep_together', 'keep_with_next',
        'left_indent', 'line_spacing', 'line_spacing_rule', 'page_break_before',
        'right_indent', 'space_after', 'space_before', 'widow_control'
    ]
    for attr in attrs:
        setattr(dest, attr, getattr(source, attr))
    for ts in source.tab_stops:
        dest.tab_stops.add_tab_stop(ts.position, ts.alignment, ts.leader)


def copyist(source, dest):
    for element in iter_block_items(source):
        if isinstance(element, Paragraph):
            new_par = dest.add_paragraph(style=element.style.name)
            copy_par_format(element.paragraph_format, new_par.paragraph_format)
            write_to_paragraph(element, new_par)
        elif isinstance(element, Table):
            new_tab = dest.add_table(len(element.rows), len(element.columns))
            new_tab.style = element.style.name
            new_tab.alignment = element.alignment
            new_tab.autofit = element.autofit
            new_tab.table_direction = element.table_direction
            for row_old, row_new in zip(element.rows, new_tab.rows):
                for cell_old, cell_new in zip(row_old.cells, row_new.cells):
                    cell_new.vertical_alignment = cell_old.vertical_alignment
                    cell_new.width = cell_old.width
                    copyist(cell_old, cell_new)


def copy_text(document: Document, destination: str):
    new = Document()
    copyist(document, new)
    new.save(destination)


def test():
    doc = Document(r'..\tab.docx')
    copy_text(doc, r'..\tabtext.docx')


def find_embeddings(document, title_first=False):
    """Yields two-element tuple with OLEObject element and paragraph element
    if OLEObject found is a ChemDraw file."""
    tree = document.element
    embedding = None
    for paragraph in tree.iterchildren():
        for run in paragraph:
            pass
    for elem in tree.iter():
        # yield only if object embedded is ChemDraw file
        if elem.tag == OLEOBJ and elem.attrib['ProgID'].startswith('ChemDraw'):
            embedding = elem
            if title_first:
                # paragraph should be already known so yield and reset
                yield (embedding, paragrph)
                embedding = None
        elif elem.tag == PAR:
            paragrph = elem
            if embedding is not None and not title_first:
                # yield and reset only if embedding found
                yield (embedding, paragrph)
                embedding = None
        else:
            continue


def extract_embeddings(
        path: pathlib.Path, out: pathlib.Path, words_in_name: int = 5,
        title_first: bool = False
             ):
    """Opens a .docx file given in `path`, reads all embedded ChemDraw
    objects and writes them to `out` directory, using first `words_in_name`
    words of the first paragraph after object (or the one directly before,
    if `title_first` is True) as the file name."""
    with zipfile.ZipFile(path) as ziphandle:
        document = Document(path)
        rels = document.part.rels
        for embedding, paragraph in find_embeddings(document, title_first):
            embedded_file = rels[embedding.attrib[RELID]].target_ref
            text = get_text(paragraph)
            title = ' '.join(text.split()[:words_in_name])
            for c in r'<>:"/\|?*':
                if c in title:
                    title = title.replace(c, "-")
                ole = ziphandle.read(embedded_file)
            with olefile.OleFileIO(ole) as ole_handle:
                content = ole_handle.openstream('CONTENTS').read()
            outfile = out / (title + '.cdx')
            try:
                with outfile.open('wb') as out_handle:
                    out_handle.write(content)
            except FileNotFoundError:
                logging.warning("Cannot write file with name: %s", title)


def get_parser():
    prs = argparse.ArgumentParser('Extract .cdx files from a .docx file.')
    prs.add_argument(
        'file', type=pathlib.Path, help='.docx file with .cdx files embedded.'
    )
    prs.add_argument(
        '-o', '--output-dir', type=pathlib.Path,
        help='Output directory. Defaults to [file-location]/[file-name]-cdx/'
    )
    prs.add_argument(
        '-w', '--words', type=int, default=5,
        help='Number of words included in name of .cdx file.'
    )
    prs.add_argument(
        '-p', '--preceding-title', action='store_true',
        help='Look for scheme description before the scheme.'
    )
    return prs


def main(argv=None):
    logging.basicConfig(level='DEBUG')
    prs = get_parser()
    args = prs.parse_args(argv)
    outdir = args.output_dir or args.file.parent / (args.file.stem + '-cdx')
    outdir.mkdir(exist_ok=True)
    extract_embeddings(args.file, outdir, args.words, args.preceding_title)


if __name__ == "__main__":
    main()
